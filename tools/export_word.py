from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
import datetime


def _get(item, key, default=""):
    if isinstance(item, dict):
        return item.get(key, default)
    return getattr(item, key, default)


def _format_extraction_stats(stats):
    stats = stats or {}
    return (
        f"Jina全文 {int(stats.get('jina_count', 0) or 0)} | "
        f"网页直连 {int(stats.get('direct_html_count', 0) or 0)} | "
        f"摘要兜底 {int(stats.get('snippet_count', 0) or 0)}"
    )


def generate_word(data, timeline_data, filename, model_name):
    doc = Document()
    normal_style = doc.styles['Normal']
    normal_style.font.name = '微软雅黑'
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    normal_style.font.size = Pt(10.5)

    for i in range(1, 4):
        h_style = doc.styles[f'Heading {i}']
        h_style.font.name = '微软雅黑'
        h_style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if i == 1:
            h_style.font.color.rgb = RGBColor(0, 51, 102)

    title = doc.add_heading("DeepSeek 企业级深度科技研报", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run(
        f"生成日期: {datetime.date.today()}  |  数据来源: Tavily 商业资讯引擎  |  分析模型: {model_name}"
    )
    meta_run.font.color.rgb = RGBColor(128, 128, 128)
    meta_run.font.size = Pt(9)
    doc.add_paragraph("━" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    if timeline_data:
        doc.add_heading("⏱️ 卷首语：核心大事件时间线", level=1)
        for t_data in timeline_data:
            doc.add_heading(f"专题: {_get(t_data, 'topic', '未命名专题')}", level=2)
            focus_tags = _get(t_data, 'focus_tags', [])
            if focus_tags:
                doc.add_paragraph(f"重点标签: {'、'.join(focus_tags[:8])}")
            extraction_stats = _get(t_data, 'extraction_stats', {})
            if extraction_stats:
                doc.add_paragraph(f"抓取概况: {_format_extraction_stats(extraction_stats)}")
            events = _get(t_data, 'events', [])
            if not events:
                doc.add_paragraph("暂无有效时间线。").font.italic = True
                continue

            for item in events:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"[{_get(item, 'date', '近期')}] ").bold = True
                event_run = p.add_run(f"{_get(item, 'event', '未命名事件')} ")
                if _get(item, 'appears_in_later_news', False):
                    event_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    badge_run = p.add_run("【后续长新闻已展开】 ")
                    badge_run.font.color.rgb = RGBColor(192, 102, 0)
                    badge_run.bold = True
                if _get(item, 'history_status', '') == 'followup':
                    history_run = p.add_run("【历史事件延续】 ")
                    history_run.font.color.rgb = RGBColor(15, 118, 110)
                    history_run.bold = True
                source_run = p.add_run(f"({_get(item, 'source', '未知来源')})")
                source_run.font.color.rgb = RGBColor(128, 128, 128)

                if _get(item, 'history_status', '') == 'followup':
                    p_hist = doc.add_paragraph()
                    p_hist.paragraph_format.left_indent = Pt(24)
                    hist_run = p_hist.add_run("历史追踪：")
                    hist_run.bold = True
                    hist_run.font.color.rgb = RGBColor(15, 118, 110)
                    p_hist.add_run(
                        f"首次记录 {_get(item, 'first_seen', '未知')}，累计追踪 {int(_get(item, 'seen_count', 1) or 1)} 次"
                    )

                if _get(item, 'appears_in_later_news', False):
                    matched_title = _get(item, 'matched_news_title', '')
                    match_reason = _get(item, 'match_reason', '')
                    if matched_title:
                        p_title = doc.add_paragraph()
                        p_title.paragraph_format.left_indent = Pt(24)
                        title_run = p_title.add_run("对应长新闻：")
                        title_run.bold = True
                        title_run.font.color.rgb = RGBColor(192, 102, 0)
                        p_title.add_run(matched_title)
                    if match_reason:
                        p_reason = doc.add_paragraph()
                        p_reason.paragraph_format.left_indent = Pt(24)
                        reason_run = p_reason.add_run("出现原因：")
                        reason_run.bold = True
                        reason_run.font.color.rgb = RGBColor(192, 102, 0)
                        p_reason.add_run(match_reason)
        doc.add_paragraph("━" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    for section in data:
        doc.add_heading(f"🔷 深度研报：{_get(section, 'topic', '未命名专题')}", level=1)
        focus_tags = _get(section, 'focus_tags', [])
        if focus_tags:
            doc.add_paragraph(f"重点标签: {'、'.join(focus_tags[:8])}")
        extraction_stats = _get(section, 'extraction_stats', {})
        if extraction_stats:
            doc.add_paragraph(f"抓取概况: {_format_extraction_stats(extraction_stats)}")
        for warning_text in _get(section, 'warnings', []):
            p_warning = doc.add_paragraph()
            warning_run = p_warning.add_run(f"注意：{warning_text}")
            warning_run.bold = True
            warning_run.font.color.rgb = RGBColor(192, 102, 0)
        news_items = _get(section, 'data', [])
        if not news_items:
            doc.add_paragraph("    在指定时间范围内，未发现符合标准的重大情报。").font.italic = True
            continue

        for news in news_items:
            doc.add_heading(f"🔹 {_get(news, 'title', '未命名情报')}", level=2)
            p_info = doc.add_paragraph()
            event_id_text = _get(news, 'event_id', '')
            run_info = p_info.add_run(
                f"    📌 来源: {_get(news, 'source', '未知来源')}    |    🕒 时间: {_get(news, 'date_check', '近期')}    |    🔥 热度: {'⭐' * int(_get(news, 'importance', 3) or 3)}"
            )
            run_info.font.color.rgb = RGBColor(100, 100, 100)
            run_info.font.bold = True
            if event_id_text:
                p_info.add_run(f"    |    🧷 事件ID: {event_id_text}").bold = True

            timeline_refs = _get(news, 'timeline_refs', [])
            if timeline_refs:
                p_linked = doc.add_paragraph()
                linked_run = p_linked.add_run("关联核心时间线：")
                linked_run.bold = True
                linked_run.font.color.rgb = RGBColor(192, 102, 0)
                for ref in timeline_refs:
                    ref_p = doc.add_paragraph()
                    ref_p.paragraph_format.left_indent = Pt(18)
                    ref_title = ref_p.add_run(f"[{_get(ref, 'date', '近期')}] {_get(ref, 'event', '未命名事件')}")
                    ref_title.bold = True
                    ref_title.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    reason = _get(ref, 'reason', '')
                    if reason:
                        ref_p.add_run(f"。{reason}")

            p_summary = doc.add_paragraph(_get(news, 'summary', '暂无详情'))
            p_summary.paragraph_format.line_spacing = 1.5
            p_summary.paragraph_format.first_line_indent = Pt(21)

            news_url = _get(news, 'url', '')
            if news_url:
                p_url = doc.add_paragraph()
                url_run = p_url.add_run(f"原文链接：{news_url}")
                url_run.font.color.rgb = RGBColor(0, 112, 192)

            divider = doc.add_paragraph("┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈")
            divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
            divider.runs[0].font.color.rgb = RGBColor(200, 200, 200)

    path = f"{filename}.docx"
    doc.save(path)
    return path
